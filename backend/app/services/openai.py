import os
import json
from io import BytesIO
import fitz
from docx import Document
from openai import OpenAI
import base64
from fastapi import UploadFile, HTTPException
from app.config import settings

client = OpenAI(api_key=settings.openai_api_key)

IMAGE_TYPES = {"image/jpeg", "image/png", "image/webp", "image/jpg"}
PDF_TYPE    = "application/pdf"
DOCX_TYPE   = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

SYSTEM_PROMPT = """You are a financial document parser for a Sri Lankan business.

Extract structured data from the invoice/receipt and classify it strictly as income or expense.

CLASSIFICATION RULES:
- "income"  -> YOU (the business) issued this invoice TO a customer.
- "expense" -> Someone else issued this invoice TO you.

Key signals for INCOME:
  Your company name in "From:", "Billed by:", "Issued by:", "Vendor:"

Key signals for EXPENSE:
  Your company name in "To:", "Billed to:", "Client:", "Sold to:"

CATEGORY examples:
  income  -> sales, service_fee, consulting, rental_income, interest
  expense -> utilities, rent, supplies, salary, transport, maintenance, tax

Return ONLY valid JSON, no markdown, no explanation:
{
  "title": "vendor or client name",
  "amount": 0.00,
  "currency": "LKR",
  "category": "category name",
  "date": "YYYY-MM-DD",
  "payment_method": "cash|card|bank_transfer|cheque or null",
  "invoice_type": "income or expense",
  "invoice_number": "INV-xxx or null",
  "tax_amount": 0.00,
  "description": "one sentence summary",
  "confidence": "high|medium|low",
  "raw_amount_text": "original amount string from document"
}"""


def _extract_pdf(raw: bytes) -> tuple[str | None, bytes | None]:
    doc = fitz.open(stream=raw, filetype="pdf")
    text = "\n".join(page.get_text() for page in doc).strip()
    if text:
        return text[:12000], None
    pix = doc[0].get_pixmap(dpi=150)
    return None, pix.tobytes("png")


def _extract_docx(raw: bytes) -> str:
    doc = Document(BytesIO(raw))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    return "\n".join(lines)[:12000]


def _parse(raw_text: str) -> dict:
    cleaned = (
        raw_text.strip()
        .removeprefix("```json")
        .removeprefix("```")
        .removesuffix("```")
        .strip()
    )
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=422, detail=f"Parse error: {e} | Raw: {cleaned[:300]}")


def call_text(text: str) -> dict:
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user",   "content": f"Extract invoice data:\n\n{text}"}
        ],
        max_tokens=1000,
        temperature=0
    )
    return _parse(response.choices[0].message.content)


def call_vision(img_bytes: bytes, media_type: str = "image/png") -> dict:
    b64_image = base64.b64encode(img_bytes).decode("utf-8")
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{media_type};base64,{b64_image}",
                            "detail": "high"
                        }
                    },
                    {
                        "type": "text",
                        "text": "Extract all invoice data from this image."
                    }
                ]
            }
        ],
        max_tokens=1000,
        temperature=0
    )
    return _parse(response.choices[0].message.content)


async def extract_invoice(file: UploadFile) -> dict:
    raw  = await file.read()
    ct   = (file.content_type or "").lower()
    name = (file.filename or "").lower()

    if ct in IMAGE_TYPES or name.endswith((".jpg", ".jpeg", ".png", ".webp")):
        return call_vision(raw, ct if ct in IMAGE_TYPES else "image/jpeg")

    if ct == PDF_TYPE or name.endswith(".pdf"):
        text, img_bytes = _extract_pdf(raw)
        return call_vision(img_bytes, "image/png") if img_bytes else call_text(text)

    if ct == DOCX_TYPE or name.endswith(".docx"):
        return call_text(_extract_docx(raw))

    return call_text(raw.decode("utf-8", errors="ignore")[:12000])